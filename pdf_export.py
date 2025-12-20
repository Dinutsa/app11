"""
Модуль експорту звіту у формат Word (.docx).
ВЕРСІЯ: DATE ON TOP (Дата зверху).
- Виправлено імпорт pandas (pd).
- Додано дату справа перед заголовком.
"""

import io
import textwrap
from datetime import datetime  # <-- Імпорт дати
import pandas as pd            # <-- ВИПРАВЛЕНО: Імпорт pandas як pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from classification import QuestionType
from summary import QuestionSummary
from typing import List

# --- НАЛАШТУВАННЯ ---
CHART_DPI = 150
FONT_SIZE_CHART = 10
BAR_WIDTH = 0.6

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def create_chart_image(qs: QuestionSummary) -> io.BytesIO:
    plt.close('all')
    plt.clf()
    plt.rcParams.update({'font.size': FONT_SIZE_CHART})
    
    labels = qs.table["Варіант відповіді"].astype(str).tolist()
    values = qs.table["Кількість"]
    wrapped_labels = [textwrap.fill(l, 25) for l in labels]

    is_scale = (qs.question.qtype == QuestionType.SCALE)
    if not is_scale:
        try:
            vals = pd.to_numeric(qs.table["Варіант відповіді"], errors='coerce')
            if vals.notna().all() and vals.min() >= 0 and vals.max() <= 10:
                is_scale = True
        except: pass

    if is_scale:
        fig = plt.figure(figsize=(6.0, 4.0))
        bars = plt.bar(wrapped_labels, values, color='#4F81BD', width=BAR_WIDTH)
        plt.ylabel('Кількість')
        plt.grid(axis='y', linestyle='--', alpha=0.5)
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                     f'{int(height)}', ha='center', va='bottom', fontweight='bold')
    else:
        fig = plt.figure(figsize=(6.0, 4.0))
        colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#4BACC6', '#F79646']
        c_arg = colors[:len(values)] if len(values) <= len(colors) else None
        wedges, texts, autotexts = plt.pie(
            values, labels=None, autopct='%1.1f%%', startangle=90,
            pctdistance=0.8, colors=c_arg, radius=1.0,
            textprops={'fontsize': FONT_SIZE_CHART}
        )
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_weight('bold')
            import matplotlib.patheffects as path_effects
            autotext.set_path_effects([path_effects.withStroke(linewidth=2, foreground='#333333')])

        plt.axis('equal')
        cols = 2 if len(labels) > 3 else 1
        plt.legend(wrapped_labels, loc="upper center", bbox_to_anchor=(0.5, 0.0), ncol=cols, frameon=False, fontsize=9)

    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=CHART_DPI, bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def build_docx_report(original_df, sliced_df, summaries, range_info) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- ДАТА СПРАВА ЗВЕРХУ ---
    date_str = datetime.now().strftime("%d.%m.%Y")
    p_date = doc.add_paragraph(date_str)
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Основний заголовок
    head = doc.add_heading('Звіт про результати опитування', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Всього анкет: {len(original_df)}")
    doc.add_paragraph(f"Оброблено: {len(sliced_df)}")
    doc.add_paragraph(f"Діапазон: {range_info}")
    doc.add_page_break()

    for qs in summaries:
        if qs.table.empty: continue
        
        p = doc.add_paragraph()
        runner = p.add_run(f"{qs.question.code}. {qs.question.text}")
        runner.bold = True
        runner.font.size = Pt(14)
        
        table = doc.add_table(rows=1, cols=3)
        set_table_borders(table)
        hdr = table.rows[0].cells
        hdr[0].text = 'Варіант'; hdr[1].text = 'Кількість'; hdr[2].text = '%'
        
        for row in qs.table.itertuples(index=False):
            rc = table.add_row().cells
            rc[0].text = str(row[0])
            rc[1].text = str(row[1])
            rc[2].text = str(row[2])

        try:
            img_stream = create_chart_image(qs)
            doc.add_picture(img_stream, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: pass
        doc.add_paragraph("\n")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()